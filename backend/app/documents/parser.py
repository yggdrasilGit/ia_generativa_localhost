from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader


@dataclass
class ParsedDocument:
    text: str
    title: str
    total_pages: int | None
    needs_ocr: bool
    page_texts: list[tuple[int, str]] | None = None


def parse_document(file_path: str) -> ParsedDocument:
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".txt":
        text = path.read_text(encoding="utf-8", errors="ignore")
        return ParsedDocument(text=text, title=path.stem, total_pages=None, needs_ocr=False, page_texts=None)

    if suffix == ".docx":
        doc = DocxDocument(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
        text = "\n".join(paragraphs)
        return ParsedDocument(text=text, title=path.stem, total_pages=None, needs_ocr=False, page_texts=None)

    if suffix == ".pdf":
        reader = PdfReader(file_path)
        pages: list[str] = []
        page_texts: list[tuple[int, str]] = []
        for idx, page in enumerate(reader.pages, start=1):
            page_text = page.extract_text() or ""
            pages.append(page_text)
            page_texts.append((idx, page_text))
        text = "\n\n".join(pages).strip()
        needs_ocr = len(text) < 80
        return ParsedDocument(
            text=text,
            title=path.stem,
            total_pages=len(reader.pages),
            needs_ocr=needs_ocr,
            page_texts=page_texts,
        )

    raise ValueError(f"Tipo de arquivo não suportado: {suffix}")
